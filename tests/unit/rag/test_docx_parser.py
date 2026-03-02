"""Unit tests for DOCXParser legal structure extraction."""

from pathlib import Path

import pytest
from docx import Document

from src.rag.parser.docx_parser import DOCXParser


def _build_legal_docx(file_path: Path) -> None:
    doc = Document()

    doc.add_paragraph("TÍTULO I - DAS DISPOSIÇÕES GERAIS")
    doc.add_paragraph("CAPÍTULO I - DAS REGRAS INICIAIS")
    doc.add_paragraph("SEÇÃO I - DA APLICAÇÃO")
    doc.add_paragraph("Art. 1o Esta lei estabelece normas gerais.")
    doc.add_paragraph("§ 1o O servidor responderá administrativamente;")
    doc.add_paragraph("nos termos do regulamento interno.")
    doc.add_paragraph("Inciso IV - Exercer com zelo e dedicação.")
    doc.add_paragraph("Nota: redação atualizada pela Lei 9.999/99.")
    doc.add_paragraph("Dispositivo revogado pela Lei 10.000/00.")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Art. 2º O quadro de pessoal será organizado por carreira."

    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "Nota de rodapé: referência histórica."

    doc.save(file_path)


@pytest.mark.unit
class TestDOCXParserLegalStructure:
    """Tests for legal structure parsing without Word heading styles."""

    def test_parse_detects_legal_blocks_without_heading_styles(self, tmp_path) -> None:
        """Should detect legal hierarchy markers from text patterns."""
        file_path = tmp_path / "lei.docx"
        _build_legal_docx(file_path)

        parsed = DOCXParser(file_path).parse()
        by_text = {item["text"]: item for item in parsed}

        titulo = by_text["TÍTULO I - DAS DISPOSIÇÕES GERAIS"]
        assert titulo["is_heading"] is True
        assert titulo["heading_level"] == 1
        assert titulo["legal_structure_kind"] == "titulo"

        capitulo = by_text["CAPÍTULO I - DAS REGRAS INICIAIS"]
        assert capitulo["is_heading"] is True
        assert capitulo["heading_level"] == 2
        assert capitulo["legal_structure_kind"] == "capitulo"

        artigo = by_text["Art. 1o Esta lei estabelece normas gerais."]
        assert artigo["is_heading"] is False
        assert artigo["legal_structure_kind"] == "artigo"
        assert artigo["legal_structure_value"] == "1"

        paragrafo = by_text["§ 1o O servidor responderá administrativamente;"]
        assert paragrafo["legal_structure_kind"] == "paragrafo"
        assert paragrafo["legal_structure_value"] == "1"

        continuation = by_text["nos termos do regulamento interno."]
        assert continuation["legal_continuation_of"] == "paragrafo"

        inciso = by_text["Inciso IV - Exercer com zelo e dedicação."]
        assert inciso["legal_structure_kind"] == "inciso"
        assert inciso["legal_structure_value"] == "IV"

    def test_parse_includes_tables_and_footers_with_legal_flags(self, tmp_path) -> None:
        """Should include table/footer paragraphs and mark notes/revocations."""
        file_path = tmp_path / "lei_tabela_rodape.docx"
        _build_legal_docx(file_path)

        parsed = DOCXParser(file_path).parse()
        by_text = {item["text"]: item for item in parsed}

        tabela = by_text["Art. 2º O quadro de pessoal será organizado por carreira."]
        assert tabela["source"].startswith("table:")
        assert tabela["legal_structure_kind"] == "artigo"
        assert tabela["legal_structure_value"] == "2"

        rodape = by_text["Nota de rodapé: referência histórica."]
        assert rodape["source"].startswith("footer:")
        assert rodape["legal_flags"]["is_nota"] is True

        nota = by_text["Nota: redação atualizada pela Lei 9.999/99."]
        assert nota["legal_flags"]["is_nota"] is True

        revogacao = by_text["Dispositivo revogado pela Lei 10.000/00."]
        assert revogacao["legal_flags"]["is_revogacao"] is True
