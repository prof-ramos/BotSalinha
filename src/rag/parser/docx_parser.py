"""DOCX document parser implementation."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any

import structlog
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.rag.utils.normalizer import normalize_encoding

log = structlog.get_logger(__name__)


class DOCXParser:
    """
    Parser for Microsoft Word (.docx) documents.

    Extracts structured text with style information, headings,
    and formatting details from DOCX files.
    """

    _LEGAL_TITLE_PATTERNS: tuple[tuple[str, re.Pattern[str], int], ...] = (
        ("titulo", re.compile(r"^\s*T[ÍI]TULO\s+([IVXLCDM]+|\d+)\b", re.IGNORECASE), 1),
        ("capitulo", re.compile(r"^\s*CAP[ÍI]TULO\s+([IVXLCDM]+|\d+)\b", re.IGNORECASE), 2),
        ("secao", re.compile(r"^\s*SE[CÇ][ÃA]O\s+([IVXLCDM]+|\d+)\b", re.IGNORECASE), 3),
        (
            "subsecao",
            re.compile(r"^\s*SUBSE[CÇ][ÃA]O\s+([IVXLCDM]+|\d+)\b", re.IGNORECASE),
            4,
        ),
    )
    _LEGAL_ARTIGO_RE = re.compile(r"^\s*Art\.?\s*([0-9]+(?:-[A-Za-z])?)(?:\s*[º°o])?\b", re.IGNORECASE)
    _LEGAL_PARAGRAFO_RE = re.compile(
        r"^\s*(?:§+\s*([0-9]+|[Uu][Nn][ÍI]?[Cc][Oo])(?:\s*[º°o])?"
        r"|Par[áa]grafo\s+([ÚUu]nico|\d+)(?:\s*[º°o])?)\b",
        re.IGNORECASE,
    )
    _LEGAL_INCISO_RE = re.compile(
        r"^\s*(?:Inciso\s+)?([IVXLCDM]{1,15})\s*(?:[-–—.)]|$)",
        re.IGNORECASE,
    )
    _LEGAL_REVOGACAO_RE = re.compile(
        r"\b(revogad[oa]|vetad[oa]|reda[cç][ãa]o\s+dada\s+pela)\b",
        re.IGNORECASE,
    )
    _LEGAL_NOTA_RE = re.compile(
        r"^\s*(nota(?:\s+de\s+rodap[eé])?|observa[cç][ãa]o|coment[aá]rio|n\.?\s*b\.?)\s*[:\-]",
        re.IGNORECASE,
    )
    _LEGAL_VETO_RE = re.compile(r"\bvetad[oa]\b|\bVETADO\b", re.IGNORECASE)
    _LEGAL_REVOGACAO_PARCIAL_RE = re.compile(r"\brevogad[oa]\s+parcialmente\b", re.IGNORECASE)
    _LEGAL_VETO_PARCIAL_RE = re.compile(r"\bvetad[oa]\s+parcialmente\b", re.IGNORECASE)

    def __init__(self, file_path: str | Path) -> None:
        """
        Initialize the DOCX parser.

        Args:
            file_path: Path to the DOCX file to parse

        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file is not a .docx file
        """
        self._file_path = Path(file_path)

        if not self._file_path.exists():
            msg = f"File not found: {file_path}"
            log.error("rag_parser_file_not_found", error=msg, file_path=str(self._file_path))
            raise FileNotFoundError(msg)

        if self._file_path.suffix.lower() != ".docx":
            msg = f"Expected .docx file, got: {self._file_path.suffix}"
            log.error("rag_parser_invalid_format", error=msg, file_path=str(self._file_path))
            raise ValueError(msg)

        log.debug(
            "rag_parser_initialized",
            file_path=str(self._file_path),
            parser="DOCXParser",
        )

    def parse(self) -> list[dict[str, Any]]:
        """
        Parse the DOCX file and extract structured content.

        Returns:
            A list of dictionaries, one per paragraph, containing:
                - text: str (texto do parágrafo)
                - style: str | None (nome do estilo)
                - is_heading: bool (se é Heading 1-9)
                - heading_level: int | None (1-9 se for heading)
                - is_bold: bool
                - is_italic: bool
                - runs: list[dict] com formatação detalhada

        Raises:
            Exception: If the document cannot be parsed
        """
        log.info(
            "rag_parser_progress",
            file_path=str(self._file_path),
            stage="started",
        )

        try:
            doc = Document(str(self._file_path))
            paragraphs_data: list[dict[str, Any]] = []
            previous_legal_kind: str | None = None
            previous_text: str = ""

            for para_idx, (paragraph, source) in enumerate(self._iter_body_blocks(doc), start=1):
                para_data = self._parse_paragraph(
                    paragraph=paragraph,
                    para_idx=para_idx,
                    source=source,
                    previous_legal_kind=previous_legal_kind,
                    previous_text=previous_text,
                )
                paragraphs_data.append(para_data)

                legal_kind = para_data.get("legal_structure_kind")
                if legal_kind:
                    previous_legal_kind = legal_kind
                previous_text = para_data.get("text", "")

            for footer_idx, paragraph in enumerate(self._iter_footer_paragraphs(doc), start=1):
                para_data = self._parse_paragraph(
                    paragraph=paragraph,
                    para_idx=len(paragraphs_data) + 1,
                    source=f"footer:{footer_idx}",
                    previous_legal_kind=previous_legal_kind,
                    previous_text=previous_text,
                )
                paragraphs_data.append(para_data)

            log.info(
                "rag_parser_progress",
                file_path=str(self._file_path),
                stage="completed",
                paragraphs_count=len(paragraphs_data),
            )

            return paragraphs_data

        except (zipfile.BadZipFile, KeyError, PackageNotFoundError) as e:
            log.error(
                "rag_parser_error",
                error=str(e),
                file_path=str(self._file_path),
            )
            raise

    def _iter_body_blocks(self, doc: Document) -> list[tuple[Paragraph, str]]:
        """Iterate body paragraphs including table cell paragraphs in document order."""
        blocks: list[tuple[Paragraph, str]] = []

        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                blocks.append((Paragraph(child, doc), "body"))
                continue

            if isinstance(child, CT_Tbl):
                table = Table(child, doc)
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            blocks.append((paragraph, f"table:r{row_idx}c{col_idx}"))

        return blocks

    def _iter_footer_paragraphs(self, doc: Document) -> list[Paragraph]:
        """Collect footer paragraphs (rodapé) across sections."""
        footer_paragraphs: list[Paragraph] = []
        seen: set[tuple[str, str]] = set()

        for section in doc.sections:
            for paragraph in section.footer.paragraphs:
                normalized = normalize_encoding(paragraph.text.strip())
                if not normalized:
                    continue
                key = (normalized, paragraph.style.name if paragraph.style else "")
                if key in seen:
                    continue
                seen.add(key)
                footer_paragraphs.append(paragraph)

        return footer_paragraphs

    def _parse_paragraph(
        self,
        paragraph: Any,
        para_idx: int,
        source: str,
        previous_legal_kind: str | None,
        previous_text: str,
    ) -> dict[str, Any]:
        """
        Parse a single paragraph and extract its properties.

        Args:
            paragraph: python-docx Paragraph object
            para_idx: Paragraph index for logging

        Returns:
            Dictionary with paragraph data
        """
        text = normalize_encoding(paragraph.text.strip())
        style_name = paragraph.style.name if paragraph.style else None

        is_heading, heading_level = self._get_heading_level(style_name)
        legal_structure = self._extract_legal_structure(text)
        revoked = bool(self._LEGAL_REVOGACAO_RE.search(text))
        vetoed = bool(self._LEGAL_VETO_RE.search(text))
        legal_flags = {
            "is_revogacao": revoked,
            "is_vetado": vetoed,
            "is_nota": bool(self._LEGAL_NOTA_RE.search(text)),
            "revocation_scope": (
                "partial"
                if self._LEGAL_REVOGACAO_PARCIAL_RE.search(text)
                else ("total" if revoked else "none")
            ),
            "veto_scope": (
                "partial"
                if self._LEGAL_VETO_PARCIAL_RE.search(text)
                else ("total" if vetoed else "none")
            ),
        }

        legal_heading_level = legal_structure.get("heading_level") if legal_structure else None
        if not is_heading and legal_heading_level is not None:
            is_heading = True
            heading_level = legal_heading_level

        # Extract runs for detailed formatting
        runs = self._extract_runs(paragraph)

        # Determine if paragraph is bold/italic (if any run is)
        is_bold = any(run.get("is_bold", False) for run in runs)
        is_italic = any(run.get("is_italic", False) for run in runs)
        continuation_of = self._detect_continuation(
            text=text,
            previous_legal_kind=previous_legal_kind,
            previous_text=previous_text,
            current_legal_kind=legal_structure.get("kind") if legal_structure else None,
        )

        return {
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "heading_level": heading_level,
            "is_bold": is_bold,
            "is_italic": is_italic,
            "runs": runs,
            "source": source,
            "paragraph_index": para_idx,
            "legal_structure_kind": legal_structure.get("kind") if legal_structure else None,
            "legal_structure_value": legal_structure.get("value") if legal_structure else None,
            "is_legal_heading": legal_heading_level is not None,
            "legal_continuation_of": continuation_of,
            "legal_flags": legal_flags,
        }

    def _extract_legal_structure(self, text: str) -> dict[str, Any] | None:
        """Detect legal block markers in plain text paragraphs."""
        if not text:
            return None

        for kind, pattern, level in self._LEGAL_TITLE_PATTERNS:
            match = pattern.search(text)
            if match:
                return {"kind": kind, "value": match.group(1), "heading_level": level}

        artigo_match = self._LEGAL_ARTIGO_RE.search(text)
        if artigo_match:
            return {
                "kind": "artigo",
                "value": self._normalize_artigo_value(artigo_match.group(1)),
                "heading_level": None,
            }

        paragrafo_match = self._LEGAL_PARAGRAFO_RE.search(text)
        if paragrafo_match:
            value = paragrafo_match.group(1) or paragrafo_match.group(2)
            return {"kind": "paragrafo", "value": value, "heading_level": None}

        inciso_match = self._LEGAL_INCISO_RE.search(text)
        if inciso_match:
            return {"kind": "inciso", "value": inciso_match.group(1).upper(), "heading_level": None}

        if self._LEGAL_NOTA_RE.search(text):
            return {"kind": "nota", "value": None, "heading_level": None}

        if self._LEGAL_REVOGACAO_RE.search(text):
            return {"kind": "revogacao", "value": None, "heading_level": None}

        return None

    @staticmethod
    def _normalize_artigo_value(value: str) -> str:
        """Normalize article number preserving forms like 10-A."""
        normalized = value.strip()
        if re.match(r"^\d+[oO]$", normalized):
            return normalized[:-1]
        return normalized

    @staticmethod
    def _detect_continuation(
        text: str,
        previous_legal_kind: str | None,
        previous_text: str,
        current_legal_kind: str | None,
    ) -> str | None:
        """Mark likely multiline continuation for legal paragraphs/incisos."""
        if not text or current_legal_kind is not None or previous_legal_kind is None:
            return None

        if previous_legal_kind not in {"paragrafo", "inciso", "artigo"}:
            return None

        continuation_starts = text[:1].islower()
        previous_opens_clause = previous_text.rstrip().endswith((",", ";", ":"))
        return previous_legal_kind if (continuation_starts or previous_opens_clause) else None

    def _extract_runs(self, paragraph: Any) -> list[dict[str, Any]]:
        """
        Extract detailed formatting information from paragraph runs.

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            List of dictionaries with run formatting details
        """
        runs_data: list[dict[str, Any]] = []

        for run in paragraph.runs:
            run_data = {
                "text": run.text,
                "is_bold": run.bold if run.bold is not None else False,
                "is_italic": run.italic if run.italic is not None else False,
                "is_underline": run.underline not in (None, False),
                "font_name": run.font.name if run.font.name else None,
                "font_size": run.font.size.pt if run.font.size else None,
            }
            runs_data.append(run_data)

        return runs_data

    def _get_heading_level(self, style_name: str | None) -> tuple[bool, int | None]:
        """
        Determine if a style is a heading and extract its level.

        Args:
            style_name: The paragraph style name

        Returns:
            Tuple of (is_heading: bool, heading_level: int | None)
        """
        if style_name is None:
            return False, None

        style_lower = style_name.lower()

        # Check for Heading 1-9 patterns (both English and Portuguese)
        # Common patterns: "Heading 1", "Título 1", "Cabeçalho 1"
        heading_patterns = [
            "heading",
            "título",
            "titulo",
            "cabeçalho",
            "cabecalho",
            "title",
            "header",
        ]

        for pattern in heading_patterns:
            if pattern in style_lower:
                # Try to extract the number
                for i in range(1, 10):  # Check 1-9
                    if f"{pattern} {i}" in style_lower or f"{pattern}{i}" in style_lower:
                        return True, i

        # Alternative: check for style names like "Heading1" (no space)
        match = re.search(
            r"(?:heading|t[ií]tulo|cabec[aá]lho|title|header)(\d)", style_lower, re.IGNORECASE
        )
        if match:
            level = int(match.group(1))
            if 1 <= level <= 9:
                return True, level

        return False, None


__all__ = ["DOCXParser"]
